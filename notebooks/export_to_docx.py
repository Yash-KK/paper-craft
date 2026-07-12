"""
Export final_paper / final_answer_key (from generation_graph.py) into styled .docx files
matching the look of the school's own sample question paper.

Usage:
    from export_to_docx import export_question_paper_and_answer_key

    export_question_paper_and_answer_key(
        question_paper=question_paper,          # the QuestionPaperBlueprint used to generate
        final_paper=result["final_paper"],
        final_answer_key=result["final_answer_key"],
        reference_docx="sample.docx",            # the school's sample QP, used as the style template
        paper_out="generated_question_paper.docx",
        answer_key_out="generated_answer_key.docx",
    )

How math is handled:
    Generated text contains LaTeX (\\(...\\), \\[...\\]). Rather than trying to hand-roll LaTeX
    rendering, each line of text is piped through `pandoc` (markdown -> docx) which converts
    $...$/$$...$$ math into native Word equations (OMML) automatically; the resulting run/
    equation XML is then spliced into the document being built, with font/size/bold re-applied
    to the surrounding text runs (equations keep Word's standard math font, which is expected -
    real exam papers render equations in a distinct font from body text too).

Requires:
    pip install python-docx pydantic
    pandoc installed on PATH (same dependency already used by blueprint_extraction.py)
"""

import copy
import io
import re
import subprocess
import zipfile
from pathlib import Path

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run

FONT_NAME = "Bookman Old Style"
SIZE_SCHOOL_NAME = 16
SIZE_EXAM_TITLE = 14
SIZE_BODY = 12

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ===========================================================================
# 1. LATEX -> NATIVE WORD EQUATIONS (via pandoc), spliced into python-docx
# ===========================================================================

MATH_SPAN_RE = re.compile(r"(\$\$.*?\$\$|\$[^$]+\$)", re.DOTALL)


def normalize_latex_delimiters(text: str) -> str:
    text = text.replace("\\[", "$$").replace("\\]", "$$")
    text = text.replace("\\(", "$").replace("\\)", "$")
    return text


def escape_markdown_outside_math(text: str) -> str:
    """Escape markdown-sensitive characters (*, _, [, ], `, \\) everywhere EXCEPT inside
    $...$/$$...$$ spans, so stray underscores/asterisks in plain prose don't get misread
    as emphasis markers by pandoc."""
    parts = MATH_SPAN_RE.split(text)
    out = []
    for part in parts:
        if part and MATH_SPAN_RE.fullmatch(part):
            out.append(part)
        else:
            out.append(re.sub(r"([\\*_\[\]`])", r"\\\1", part))
    return "".join(out)


def prepare_markdown(text: str) -> str:
    return escape_markdown_outside_math(normalize_latex_delimiters(text))


def _pandoc_markdown_to_document_xml(md_text: str) -> bytes:
    result = subprocess.run(
        ["pandoc", "-f", "markdown", "-t", "docx", "-o", "-"],
        input=md_text.encode("utf-8"), capture_output=True, check=True,
    )
    with zipfile.ZipFile(io.BytesIO(result.stdout)) as z:
        return z.read("word/document.xml")


def _extract_first_paragraph_children(document_xml: bytes):
    root = parse_xml(document_xml)
    body = root.find(f"{{{W_NS}}}body")
    p = body.find(f"{{{W_NS}}}p")
    if p is None:
        return []
    # Skip pandoc's own <w:pPr> (paragraph style/props) - we only want the actual
    # content (runs, math). Keeping it would create a second, invalid <w:pPr> in our
    # paragraph alongside the one we set ourselves (e.g. alignment).
    return [child for child in p if etree_localname(child) != "pPr"]


def etree_localname(element) -> str:
    from lxml import etree
    return etree.QName(element).localname


def _tighten_spacing(paragraph):
    """sample.docx sets space_after=0 on every single paragraph (tightly packed layout).
    Without this, paragraphs fall back to Word's default ~8pt-after spacing and the
    document balloons into a sprawling, airy-looking mess across many extra pages."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    return paragraph


def add_blank_line(doc):
    """A deliberate, controlled spacer - use this instead of a bare doc.add_paragraph()
    so spacing stays consistent/tight everywhere else."""
    return _tighten_spacing(doc.add_paragraph())


def add_rich_paragraph(doc, text, bold=False, italic=False, size_pt=SIZE_BODY,
                        alignment=None, font_name=FONT_NAME, indent_cm=None):
    """Adds one paragraph, converting any $...$/$$...$$ math into a native Word equation.
    `text` should be a single logical line - split multi-line content before calling this
    (see add_rich_block below) so each paragraph/line break in the source shows up as a
    real paragraph break in the document."""
    paragraph = doc.add_paragraph()
    _tighten_spacing(paragraph)
    if alignment is not None:
        paragraph.alignment = alignment
    if indent_cm is not None:
        paragraph.paragraph_format.left_indent = docx.shared.Cm(indent_cm)

    if not text.strip():
        return paragraph

    md = prepare_markdown(text)
    children = _extract_first_paragraph_children(_pandoc_markdown_to_document_xml(md))
    for child in children:
        from lxml import etree
        tag = etree.QName(child).localname
        el = copy.deepcopy(child)
        if tag == "r":
            run_obj = Run(el, paragraph)
            run_obj.font.name = font_name
            run_obj.font.size = Pt(size_pt)
            run_obj.bold = bold
            run_obj.italic = italic
        paragraph._p.append(el)
    return paragraph


def add_rich_block(doc, text, **kwargs):
    """Splits text on any run of newlines into separate paragraphs (case-study passages,
    multi-line answers, etc. commonly come back as one string with embedded \\n\\n between
    labeled parts) and renders each as its own add_rich_paragraph call."""
    if not text:
        return
    for line in re.split(r"\n+", text):
        line = line.strip()
        if line:
            add_rich_paragraph(doc, line, **kwargs)


# ===========================================================================
# 2. TEMPLATE HANDLING (clone the school's sample.docx for fonts/margins/page setup)
# ===========================================================================

def load_template(reference_docx: str | Path) -> docx.Document:
    """Opens the school's sample QP and clears its body content while keeping styles,
    margins, page size, and any header/footer - so the generated file inherits the same
    page setup as the original."""
    doc = docx.Document(str(reference_docx))
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    return doc


def _insert_before_sectpr(doc, element):
    sect_pr = doc.element.body.find(qn("w:sectPr"))
    sect_pr.addprevious(element)


def clone_paragraphs(reference_docx: str | Path, doc, start: int, end: int):
    """Deep-copies paragraphs [start:end) from the ORIGINAL reference doc verbatim (exact
    runs, bold, fonts) into the target doc. Used for boilerplate that should be byte-for-byte
    identical to the source template (school name/address/title) rather than re-typeset."""
    source_doc = docx.Document(str(reference_docx))
    for p in source_doc.paragraphs[start:end]:
        _insert_before_sectpr(doc, copy.deepcopy(p._p))


def find_paragraph_index(reference_docx: str | Path, predicate) -> int | None:
    """Returns the index of the first paragraph in the reference doc matching predicate(text),
    or None if not found."""
    source_doc = docx.Document(str(reference_docx))
    for i, p in enumerate(source_doc.paragraphs):
        if predicate(p.text):
            return i
    return None


def add_bottom_border(paragraph, size=12, color="000000"):
    """Adds a single bottom border line to a paragraph (matches the divider under the
    school's info block - a real w:pBdr element, not a fake underline)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = parse_xml(
        f'<w:pBdr xmlns:w="{W_NS}">'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f"</w:pBdr>"
    )
    p_pr.append(p_bdr)


def set_list_numbering(doc, paragraph, num_id: int, ilvl: int = 0):
    """Attaches the template's own numbering definition (numId) to a paragraph, so
    generated instructions get real Word auto-numbering (1. 2. 3...) with correct
    hanging indent - matching the template exactly instead of literal '1. ' text."""
    paragraph.style = doc.styles["List Paragraph"]
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = parse_xml(
        f'<w:numPr xmlns:w="{W_NS}">'
        f'<w:ilvl w:val="{ilvl}"/><w:numId w:val="{num_id}"/>'
        f"</w:numPr>"
    )
    p_pr.append(num_pr)


def format_option_label(raw_option: str, index: int) -> str:
    """Normalizes any incoming option format ('a) text', '(a) text', 'a. text', or no
    label at all) to a consistent '(a) text' style, matching the sample paper's convention."""
    letter = "abcd"[index] if index < 4 else chr(ord("a") + index)
    stripped = re.sub(r"^\s*\(?[a-dA-D]\)?[.\)]\s*", "", raw_option).strip()
    return f"({letter}) {stripped}"


# ===========================================================================
# 3. HEADER BLOCK (school name, exam title, info line, general instructions)
# ===========================================================================

def header_from_question_paper(question_paper) -> dict:
    """Accepts either a QuestionPaperBlueprint instance or an equivalent dict."""
    get = (lambda k: getattr(question_paper, k)) if hasattr(question_paper, "subject") \
        else (lambda k: question_paper.get(k))
    return {
        "school_name": get("school_name"),
        "exam_title": get("exam_title"),
        "subject": get("subject"),
        "grade": get("grade"),
        "total_marks": get("total_marks"),
        "duration_minutes": get("duration_minutes"),
        "general_instructions": get("general_instructions") or [],
    }


def add_header_block(doc, header: dict, reference_docx: str | Path, title_suffix: str | None = None):
    # Clone the school name / address / exam title block VERBATIM from the template up to
    # (not including) the "NAME:" info line - this is pure boilerplate, byte-identical to
    # the source rather than re-typeset, so its exact formatting is never at risk of drifting.
    name_line_idx = find_paragraph_index(reference_docx, lambda t: t.strip().upper().startswith("NAME:"))
    if name_line_idx is not None:
        clone_paragraphs(reference_docx, doc, 0, name_line_idx)
        if title_suffix:
            add_rich_paragraph(doc, title_suffix, bold=True, size_pt=SIZE_EXAM_TITLE,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        # Fallback for a differently-structured reference doc with no "NAME:" line.
        if header.get("school_name"):
            add_rich_paragraph(doc, header["school_name"], bold=True, size_pt=SIZE_SCHOOL_NAME,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)
        title = header.get("exam_title") or ""
        if title_suffix:
            title = f"{title} - {title_suffix}" if title else title_suffix
        if title:
            add_rich_paragraph(doc, title, bold=True, size_pt=SIZE_EXAM_TITLE,
                                alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Info line: NAME/ROLL NO (blank for the student to fill in) + GRADE, then
    # SUBJECT/DATE and DURATION/MAX MARKS with the same bottom-border divider the
    # template uses. Fixed spacing (not \t) - raw tabs with no custom tab stops defined
    # jump unpredictably at this font/size and were wrapping the line onto two lines.
    grade_line = f"NAME: {'…' * 18}   ROLL NO: {'…' * 6}   GRADE: {header.get('grade', '')}"
    add_rich_paragraph(doc, grade_line, bold=True, size_pt=SIZE_BODY)

    subject_line = add_rich_paragraph(
        doc, f"SUBJECT: {header.get('subject', '')}", bold=True, size_pt=SIZE_BODY,
    )
    add_bottom_border(subject_line)

    duration_txt = ""
    if header.get("duration_minutes"):
        hours = header["duration_minutes"] / 60
        duration_txt = f"DURATION: {hours:g} HOUR(S)"
    marks_line = add_rich_paragraph(
        doc, f"{duration_txt}          MAX MARKS: {header.get('total_marks', '')}",
        bold=True, size_pt=SIZE_BODY,
    )
    add_bottom_border(marks_line)

    if header.get("general_instructions"):
        add_rich_paragraph(doc, "General Instructions:", bold=True, size_pt=SIZE_BODY)
        add_rich_paragraph(doc, "Read the following instructions carefully and strictly follow them:",
                            bold=True, size_pt=SIZE_BODY)
        for instr in header["general_instructions"]:
            p = add_rich_paragraph(doc, instr, size_pt=SIZE_BODY)
            set_list_numbering(doc, p, num_id=1)
        add_blank_line(doc)


# ===========================================================================
# 4. QUESTION PAPER (student-facing)
# ===========================================================================

def add_section_header(doc, section_name: str, section_questions: list[dict]):
    add_blank_line(doc)
    m = re.match(r"^Section\s+([A-Za-z0-9]+)$", section_name.strip(), flags=re.IGNORECASE)
    display_name = f"SECTION \u2013 {m.group(1).upper()}" if m else section_name.upper()
    add_rich_paragraph(doc, display_name, bold=True, size_pt=SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    total = sum(q["marks"] for q in section_questions)
    marks_values = {q["marks"] for q in section_questions}
    if len(marks_values) == 1:
        each = marks_values.pop()
        scheme = f"{len(section_questions)}X{each:g}={total:g}M"
    else:
        scheme = f"{len(section_questions)} questions, {total:g}M total"
    add_rich_paragraph(doc, scheme, bold=True, size_pt=SIZE_BODY, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def add_question(doc, q: dict):
    add_rich_paragraph(doc, f"{q['question_number']}. {q['question_text']}", size_pt=SIZE_BODY)

    if q.get("options"):
        formatted = [format_option_label(opt, i) for i, opt in enumerate(q["options"])]
        options_line = "    " + "      ".join(formatted)
        add_rich_paragraph(doc, options_line, size_pt=SIZE_BODY)

    if q.get("alternate_question_text"):
        add_rich_paragraph(doc, "OR", bold=True, size_pt=SIZE_BODY,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_rich_block(doc, q["alternate_question_text"], size_pt=SIZE_BODY)

    add_blank_line(doc)


def build_question_paper_docx(question_paper, final_paper: dict, reference_docx: str | Path,
                               out_path: str | Path):
    doc = load_template(reference_docx)
    add_header_block(doc, header_from_question_paper(question_paper), reference_docx)

    for section_name, questions in final_paper["sections"].items():
        add_section_header(doc, section_name, questions)
        for q in questions:
            add_question(doc, q)

    doc.save(str(out_path))
    return out_path


# ===========================================================================
# 5. ANSWER KEY (teacher-facing)
# ===========================================================================

def add_answer_item(doc, item: dict):
    label = f"Q{item['question_number']}"
    meta_bits = [b for b in [item.get("chapter_name"), item.get("blooms_level")] if b]
    if meta_bits:
        label += "  (" + ", ".join(meta_bits) + ")"
    add_rich_paragraph(doc, label, bold=True, size_pt=SIZE_BODY)

    if item.get("status") and item["status"] != "ok":
        add_rich_paragraph(doc, f"[FLAGGED FOR REVIEW: {item['status']}]", bold=True, size_pt=SIZE_BODY)

    if item.get("correct_option"):
        add_rich_paragraph(doc, f"Correct option: ({item['correct_option']})", size_pt=SIZE_BODY)

    add_rich_paragraph(doc, "Answer:", italic=True, size_pt=SIZE_BODY)
    add_rich_block(doc, item.get("answer") or "", size_pt=SIZE_BODY, indent_cm=1)

    if item.get("marking_rubric"):
        add_rich_paragraph(doc, "Marking scheme:", italic=True, size_pt=SIZE_BODY)
        for step in item["marking_rubric"]:
            add_rich_paragraph(doc, f"- {step['description']} ({step['marks']:g} marks)",
                                size_pt=SIZE_BODY, indent_cm=1)

    if item.get("alternate_answer"):
        add_rich_paragraph(doc, "OR (alternate):", italic=True, size_pt=SIZE_BODY,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_rich_block(doc, item["alternate_answer"], size_pt=SIZE_BODY, indent_cm=1)

    add_blank_line(doc)


def build_answer_key_docx(question_paper, final_answer_key: dict, reference_docx: str | Path,
                           out_path: str | Path):
    doc = load_template(reference_docx)
    add_header_block(doc, header_from_question_paper(question_paper), reference_docx, title_suffix="Answer Key")

    for section_name, items in final_answer_key["sections"].items():
        add_blank_line(doc)
        add_rich_paragraph(doc, section_name.upper(), bold=True, size_pt=SIZE_BODY,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for item in items:
            add_answer_item(doc, item)

    doc.save(str(out_path))
    return out_path


# ===========================================================================
# 6. TOP-LEVEL CONVENIENCE
# ===========================================================================

def export_question_paper_only(
    question_paper, final_paper: dict, reference_docx: str | Path,
    paper_out: str | Path = "generated_question_paper.docx",
):
    """Builds ONLY the student-facing question paper - no answers, no answer key file at
    all. This is all build_question_paper_docx ever uses: final_paper only, never
    final_answer_key - answers were never present in the question paper output."""
    return build_question_paper_docx(question_paper, final_paper, reference_docx, paper_out)


def export_question_paper_and_answer_key(
    question_paper, final_paper: dict, final_answer_key: dict,
    reference_docx: str | Path,
    paper_out: str | Path = "generated_question_paper.docx",
    answer_key_out: str | Path = "generated_answer_key.docx",
):
    build_question_paper_docx(question_paper, final_paper, reference_docx, paper_out)
    build_answer_key_docx(question_paper, final_answer_key, reference_docx, answer_key_out)
    return str(paper_out), str(answer_key_out)